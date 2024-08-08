"""
Objective: a program to automatically turn large word document (~100 pages) 
    with images into a compressed latex document zipped folder that can be 
    directly uploaded to overleaf. 
    I want all the origial images from the word doc to be saved and organized 
    in the output folder. 

Input: Large Word Doc to be converted to Latex 

Output: Compressed folder to be directly uploaded to latex overleaf. 

Author: Yujing Zou, August, 2024

"""

import os
import sys
import logging
from docx import Document
from pylatex import Document as LatexDocument, Section, Subsection, Command, Package, Figure, Table, Tabular
from pylatex.utils import NoEscape
from pylatexenc.latexencode import utf8tolatex
import zipfile
from tqdm import tqdm

# Set up logging
logging.basicConfig(filename='conversion_errors.log', level=logging.ERROR, 
                    format='%(asctime)s:%(levelname)s:%(message)s')

class DocxToLatexConverter:
    def __init__(self, docx_path, output_dir):
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"The DOCX file was not found at the specified path: {docx_path}")
        
        self.docx_path = docx_path
        self.output_dir = output_dir
        self.latex_dir = os.path.join(output_dir, 'latex')
        self.images_dir = os.path.join(self.latex_dir, 'images')
        self.doc = Document(docx_path)
        self.latex_doc = LatexDocument(documentclass='article')
        self.image_count = 0
    
    def extract_images(self):
        if not os.path.exists(self.images_dir):
            os.makedirs(self.images_dir)
        for i, rel in enumerate(self.doc.part.rels):
            try:
                if "image" in self.doc.part.rels[rel].target_ref:
                    img = self.doc.part.rels[rel].target_part.blob
                    img_name = f"image{i}.png"
                    with open(os.path.join(self.images_dir, img_name), "wb") as f:
                        f.write(img)
            except Exception as e:
                logging.error(f"Failed to extract image {i}: {e}")
        self.image_count = len(os.listdir(self.images_dir))
    
    def handle_paragraph(self, paragraph):
        latex_text = utf8tolatex(paragraph.text)
        if paragraph.style.name.startswith('Heading'):
            level = int(paragraph.style.name.split()[-1])
            if level == 1:
                self.latex_doc.append(Section(latex_text))
            else:
                self.latex_doc.append(Subsection(latex_text))
        else:
            self.latex_doc.append(latex_text)
            self.latex_doc.append('\n\n')
    
    def handle_table(self, table):
        num_cols = len(table.columns)
        tab = Tabular('|' + 'c|' * num_cols)
        tab.add_hline()
        for row in table.rows:
            tab.add_row([utf8tolatex(cell.text) for cell in row.cells])
            tab.add_hline()
        with self.latex_doc.create(Table(position='h!')) as _table:
            _table.add_caption('Table')
            _table.append(tab)
    
    def handle_equation(self, equation):
        self.latex_doc.append(NoEscape(r'\[' + utf8tolatex(equation.text) + r'\]'))
    
    def convert(self):
        os.makedirs(self.images_dir, exist_ok=True)
        self.extract_images()
        
        self.latex_doc.packages.append(Package('graphicx'))
        
        paragraphs = self.doc.paragraphs
        with tqdm(total=len(paragraphs), desc="Processing paragraphs") as pbar:
            for para in paragraphs:
                self.handle_paragraph(para)
                pbar.update(1)
        
        with tqdm(total=len(self.doc.tables), desc="Processing tables") as pbar:
            for table in self.doc.tables:
                self.handle_table(table)
                pbar.update(1)
        
        existing_images = set(os.listdir(self.images_dir))
        with tqdm(total=self.image_count, desc="Processing images") as pbar:
            for i in range(self.image_count):
                img_name = f"image{i}.png"
                if img_name in existing_images:
                    with self.latex_doc.create(Figure(position='h!')) as fig:
                        fig.add_image(f'images/{img_name}', width=NoEscape(r'0.8\textwidth'))
                        fig.add_caption(f'Image {i+1}')
                else:
                    logging.error(f"Image {i} (expected {img_name}) is missing.")
                pbar.update(1)
        
        with tqdm(total=len(paragraphs), desc="Processing equations") as pbar:
            for para in paragraphs:
                if para.style.name == 'Equation':
                    self.handle_equation(para)
                pbar.update(1)
        
        tex_path = os.path.join(self.latex_dir, 'document.tex')
        with open(tex_path, 'w') as f:
            f.write(self.latex_doc.dumps())
        
        self.create_zip()
    
    def create_zip(self):
        zip_path = os.path.join(self.output_dir, 'latex_project.zip')
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for root, _, files in os.walk(self.latex_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, self.latex_dir)
                    zipf.write(file_path, arcname)
        print(f"LaTeX project zip file created at: {zip_path}")

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python main.py <path_to_docx_file> <output_directory>")
        sys.exit(1)

    docx_path = sys.argv[1]
    output_dir = sys.argv[2]

    # Create the output directory if it does not exist
    os.makedirs(output_dir, exist_ok=True)

    try:
        converter = DocxToLatexConverter(docx_path, output_dir)
        converter.convert()
    except Exception as e:
        print(f"An error occurred: {e}")
        sys.exit(1)

# For use 
# python docx_to_latex_converter.py <path_to_docx_file> <output_directory>

# Example
# python docx_to_latex_converter.py /path/to/your/document.docx /path/to/output/directory
